from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .serializers import RegisterSerializer, LoginSerializer, UserSerializer, DocumentSerializer
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.permissions import IsAuthenticated
from .models import Profile
import docx

import os
import openai
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from .models import Document
import PyPDF2
from docx import Document as DocxDocument
from docify.settings import OPENAI_API_KEY


def get_tokens(user):
    refresh = RefreshToken.for_user(user)
    return {
        'refresh': str(refresh),
        'access': str(refresh.access_token),
    }

class RegisterView(APIView):
    permission_classes = []  # Ensure only authenticated users can upload documents

    def post(self, request):
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            profile = Profile.objects.create(user=user) 
            tokens = get_tokens(user)
            return Response({
                'message': 'User created successfully',
                'tokens': tokens
            }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class LoginView(APIView):
    permission_classes = []  # Ensure only authenticated users can upload documents

    def post(self, request):
        serializer = LoginSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.validated_data['user']
            tokens = get_tokens(user)
            return Response({
                'message': 'Login successful',
                'tokens': tokens
            }, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class DocumentUploadView(APIView):
    permission_classes = [IsAuthenticated]  # Ensure only authenticated users can upload documents

    def post(self, request):
        # Assign the logged-in user's profile
        profile = request.user.profile
        
        # Create the data dictionary with the user's profile
        data = request.data.copy()
        data['profile'] = profile.id
        
        # Initialize the serializer with the data
        serializer = DocumentSerializer(data=data)
        
        if serializer.is_valid():
            document = serializer.save()  # Save the document to the database
            return Response({
                'message': 'Document uploaded successfully.',
                'document': {
                    'title': document.title,
                    'file': document.file.url,
                    'uploaded_at': document.uploaded_at
                }
            }, status=status.HTTP_201_CREATED)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    



client = openai.OpenAI(api_key=OPENAI_API_KEY)
openai.api_key = OPENAI_API_KEY


# def extract_text(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".pdf":
#         with open(file_path, 'rb') as f:
#             reader = PyPDF2.PdfReader(f)
#             return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
#     elif ext == ".docx":
#         with open(file_path, 'rb') as f:
#             doc = DocxDocument(f)
#             return "\n".join(p.text for p in doc.paragraphs)
#     return ""

class DocChatView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        prompt = request.data.get('prompt')
        if not prompt:
            return Response({"error": "Prompt is required."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            profile = request.user.profile
            document = profile.documents.last()  
            if not document:
                return Response({"error": "No document found for this user."}, status=status.HTTP_404_NOT_FOUND)

            doc_path = document.file.path
            file_extension = os.path.splitext(doc_path)[1].lower()

            if file_extension == '.pdf':
                with open(doc_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    doc_content = ""
                    for page in reader.pages:
                        doc_content += page.extract_text()
            elif file_extension in ['.docx', '.doc']:
                doc = docx.Document(doc_path)
                doc_content = '\n'.join([para.text for para in doc.paragraphs])
            else:
                return Response({"error": "Unsupported file type."}, status=status.HTTP_400_BAD_REQUEST)

            messages = [
                {"role": "system", "content": "Determine whether the following sentence is a PROMPT — a request or instruction for the AI or system to perform an action."},
                # {"role": "user", "content": f"Document content:\n{doc_content[:3000]}"},
                {"role": "user", "content": f"Is the following text a prompt?: {prompt}"}
            ]

            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=messages
            )
            reply = response.choices[0].message.content
            return Response({"answer": reply}, status=status.HTTP_200_OK)

            messages = [
                {"role": "system", "content": "You are a document assistant. Use the provided document to answer questions."},
                {"role": "user", "content": f"Document content:\n{doc_content[:3000]}"},
                {"role": "user", "content": f"My question is: {prompt}"}
            ]

            # Get OpenAI response
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            reply = response.choices[0].message.content
            return Response({"answer": reply}, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)